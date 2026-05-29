from pathlib import Path
from dataclasses import dataclass, field
from typing import List

import fitz
from docx import Document
from bs4 import BeautifulSoup


@dataclass
class ContentBlock:
    block_type: str
    text: str
    metadata: dict = field(default_factory=dict)


def load_pdf(filepath: str, source_doc: str) -> List[ContentBlock]:
    blocks = []
    doc = fitz.open(filepath)
    for page_num, page in enumerate(doc, start=1):
        page_dict = page.get_text("dict")
        current_section = "Introduction"
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            block_text = ""
            max_font_size = 0.0
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    block_text += span.get("text", "") + " "
                    max_font_size = max(max_font_size, span.get("size", 0.0))
            block_text = block_text.strip()
            if not block_text or len(block_text) < 10:
                continue
            if max_font_size >= 16:
                block_type = "heading"
                current_section = block_text
            elif max_font_size >= 13:
                block_type = "heading"
            else:
                block_type = "paragraph"
            blocks.append(
                ContentBlock(
                    block_type=block_type,
                    text=block_text,
                    metadata={
                        "source_doc": source_doc,
                        "page_number": page_num,
                        "section": current_section,
                        "file_type": "pdf",
                    },
                )
            )
    doc.close()
    return blocks


def load_docx(filepath: str, source_doc: str) -> List[ContentBlock]:
    blocks = []
    document = Document(filepath)
    current_section = "Introduction"

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower()
        if "heading 1" in style_name or "heading 2" in style_name:
            block_type = "heading"
            current_section = text
        elif "heading" in style_name:
            block_type = "heading"
        elif "list" in style_name:
            block_type = "list_item"
        else:
            block_type = "paragraph"
        blocks.append(
            ContentBlock(
                block_type=block_type,
                text=text,
                metadata={
                    "source_doc": source_doc,
                    "page_number": 1,
                    "section": current_section,
                    "file_type": "docx",
                },
            )
        )

    for table in document.tables:
        rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                rows.append(row_text)
        if rows:
            blocks.append(
                ContentBlock(
                    block_type="table",
                    text="\n".join(rows),
                    metadata={
                        "source_doc": source_doc,
                        "page_number": 1,
                        "section": current_section,
                        "file_type": "docx",
                    },
                )
            )

    return blocks


def load_html(html_content: str, source_doc: str, source_url: str = "") -> List[ContentBlock]:
    blocks = []
    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    current_section = "Introduction"
    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
        text = element.get_text(separator=" ", strip=True)
        if not text or len(text) < 15:
            continue
        tag_name = element.name.lower()
        if tag_name in ("h1", "h2"):
            block_type = "heading"
            current_section = text
        elif tag_name in ("h3", "h4"):
            block_type = "heading"
        elif tag_name == "li":
            block_type = "list_item"
        elif tag_name == "table":
            rows = []
            for tr in element.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    rows.append(row_text)
            text = "\n".join(rows)
            block_type = "table"
        else:
            block_type = "paragraph"
        blocks.append(
            ContentBlock(
                block_type=block_type,
                text=text,
                metadata={
                    "source_doc": source_doc,
                    "source_url": source_url,
                    "page_number": 1,
                    "section": current_section,
                    "file_type": "html",
                },
            )
        )
    return blocks


def load_document(filepath: str) -> List[ContentBlock]:
    path = Path(filepath)
    source_doc = path.stem
    ext = path.suffix.lower()
    if ext == ".pdf":
        return load_pdf(filepath, source_doc)
    if ext in (".docx", ".doc"):
        return load_docx(filepath, source_doc)
    if ext in (".html", ".htm"):
        return load_html(path.read_text(encoding="utf-8"), source_doc, filepath)
    raise ValueError(f"Unsupported file type: {ext}. Supported: pdf, docx, html")
